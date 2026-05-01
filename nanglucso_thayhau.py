import streamlit as st
import io
import os
import re
from docx import Document

# 1. GIAO DIỆN CHUẨN TRẮNG - ĐEN
st.set_page_config(page_title="AI Tích Hợp NLS - Thầy Hậu", layout="wide")
st.markdown("""
    <style>
    .stApp { background-color: #FFFFFF !important; }
    h1, h2, h3, h4, p, span, div, label, li, .stMarkdown { color: #000000 !important; }
    .stButton>button { background-color: #1E40AF !important; color: white !important; font-weight: bold !important; border-radius: 8px; }
    .status-box { padding: 15px; border-radius: 8px; border-left: 5px solid #1E40AF; background-color: #F8F9FA; margin-bottom: 20px;}
    </style>
    """, unsafe_allow_html=True)

st.title("🤖 TRỢ LÝ AI: TÍCH HỢP NĂNG LỰC SỐ (BẢN ĐỌC PHỤ LỤC 3)")
st.info("💡 Tính năng mới: AI tự động đọc Cột 3 (Tên bài) & Cột 7 (NLS) từ Phụ lục 3 để gán vào giáo án một cách linh hoạt, chính xác 100%.")

col1, col2 = st.columns(2)
with col1:
    st.subheader("1. Dữ liệu nguồn")
    file_phuluc = st.file_uploader("Tải Phụ lục 3 (Bảng mã NLS)", type=["docx"])
    file_khung = st.file_uploader("Tải Khung NLS (Chi tiết)", type=["docx"])
with col2:
    st.subheader("2. Giáo án cần xử lý")
    file_giaolan = st.file_uploader("Tải Giáo án gốc (File Word)", type=["docx"])

if st.button("🚀 BẮT ĐẦU PHÂN TÍCH VÀ TÍCH HỢP"):
    if file_giaolan and file_phuluc:
        with st.spinner("AI đang đối chiếu Giáo án với Cột 3 và Cột 7 của Phụ lục 3..."):
            doc = Document(file_giaolan)
            pl_doc = Document(file_phuluc)
            
            # Đặt tên file: Tên cũ + NLS
            original_name = os.path.splitext(file_giaolan.name)[0]
            new_filename = f"{original_name} NLS.docx"

            # --- MODULE TỰ ĐỘNG LẤY DỮ LIỆU TỪ PHỤ LỤC 3 ---
            bai_hoc = ""
            # 1. Tìm tên bài trong giáo án (Dò 30 dòng đầu tiên)
            for p in doc.paragraphs[:30]:
                match = re.search(r'(bài\s+\d+)', p.text, re.IGNORECASE)
                if match:
                    bai_hoc = match.group(1).lower()
                    break
            
            nls_list = []
            found_nls_text = ""

            if bai_hoc:
                # 2. Tìm bài đó trong CỘT 3 của Phụ lục 3
                for table in pl_doc.tables:
                    for row in table.rows:
                        if len(row.cells) >= 7: # Đảm bảo bảng có ít nhất 7 cột
                            col3_text = row.cells[2].text.lower() # Cột 3 (Index 2)
                            if bai_hoc in col3_text:
                                # 3. Lấy dữ liệu ở CỘT 7 (Index 6)
                                found_nls_text += row.cells[6].text + "\n"

                # Hỗ trợ dò trong text thường (Nếu phụ lục lưu dạng CSV text)
                if not found_nls_text:
                    for p in pl_doc.paragraphs:
                        if bai_hoc in p.text.lower() and re.search(r'\d+\.\d+\.[A-Za-z0-9]+', p.text):
                            found_nls_text += p.text + "\n"

                # 4. Tách mã NLS và nội dung
                if found_nls_text:
                    pattern = r'(\d+\.\d+\.[A-Za-z0-9]+)\s*[:\-]\s*(.*?)(?=\d+\.\d+\.[A-Za-z0-9]+\s*[:\-]|$)'
                    matches = re.findall(pattern, found_nls_text, re.DOTALL)
                    for ma, nd in matches:
                        nd_clean = nd.strip().strip('"').strip()
                        if not any(item['ma'] == ma.strip() for item in nls_list):
                            nls_list.append({"ma": ma.strip(), "nd": nd_clean})
            
            # Kiểm tra xem AI có lấy được dữ liệu thật không
            if not nls_list:
                st.warning(f"⚠️ Không tìm thấy '{bai_hoc.upper()}' trong Cột 3 của Phụ lục 3, hoặc Cột 7 bị trống. AI sẽ dùng dữ liệu giả định để tiếp tục.")
                nls_list = [
                    {"ma": "5.1.TC1a", "nd": "Hướng dẫn HS thực hành thao tác thiết bị, bật/tắt đúng quy trình."},
                    {"ma": "5.2.TC1a", "nd": "Yêu cầu HS xác định nhu cầu thực tế cần giải quyết."}
                ]
            else:
                st.success(f"🎯 AI đã đối chiếu thành công **{bai_hoc.upper()}** và lấy được {len(nls_list)} mã từ Cột 7 của Phụ lục 3.")

            # --- BƯỚC 1: XỬ LÝ MỤC TIÊU (2.3 hoặc 2.4) ---
            has_khac = any("2.3." in p.text and "khác" in p.text.lower() for p in doc.paragraphs)
            num_nls = "2.4." if has_khac else "2.3."

            target_idx = -1
            for i, p in enumerate(doc.paragraphs):
                if "2.2." in p.text and "tin học" in p.text.lower():
                    target_idx = i
                    break
            
            if target_idx != -1:
                insert_pos = target_idx + 1
                for j in range(target_idx + 1, len(doc.paragraphs)):
                    t = doc.paragraphs[j].text.strip()
                    if t.startswith("3.") or t.startswith("II.") or t.startswith("III."):
                        insert_pos = j
                        break
                
                # Cố định vị trí để chèn không bị lộn ngược
                ref_p = doc.paragraphs[insert_pos]
                
                p_head = ref_p.insert_paragraph_before(f"{num_nls} Năng lực số:")
                p_head.runs[0].bold = True
                
                # Chèn danh sách NLS chuẩn form Phụ lục 3
                for item in nls_list:
                    p_item = ref_p.insert_paragraph_before(f"- {item['ma']}: {item['nd']}")
                    p_item.runs[0].italic = True

            # --- BƯỚC 2: RÀ SOÁT TÌM "CHUYỂN GIAO NHIỆM VỤ" TRONG TIẾN TRÌNH ---
            tracker = {"idx": 0, "count": 0}

            def inject_nls(para):
                txt = para.text.lower()
                if "chuyển giao nhiệm vụ" in txt or "giao nhiệm vụ:" in txt:
                    if len(nls_list) > 0:
                        nls = nls_list[tracker["idx"] % len(nls_list)]
                        
                        # Chèn thêm 1 dòng mới vào dưới "Chuyển giao nhiệm vụ"
                        run = para.add_run(f"\n   => Tích hợp NLS: {nls['ma']}: {nls['nd']}")
                        
                        # Định dạng chuẩn: Đậm, Nghiêng, Gạch chân
                        run.italic = True
                        run.underline = True
                        run.bold = True 
                        
                        tracker["idx"] += 1
                        tracker["count"] += 1

            # Quét ngoài bảng
            for p in doc.paragraphs:
                inject_nls(p)
                
            # Quét sâu trong các bảng
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            inject_nls(p)

            # XUẤT FILE
            bio = io.BytesIO()
            doc.save(bio)
            
            if tracker["count"] > 0:
                st.markdown(f"<div class='status-box'><b>✅ BÁO CÁO NHANH:</b> Đã chèn {tracker['count']} năng lực số vào các bước 'Chuyển giao nhiệm vụ' thành công!</div>", unsafe_allow_html=True)
            
            st.download_button(
                label=f"📥 TẢI XUỐNG: {new_filename}",
                data=bio.getvalue(),
                file_name=new_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("⚠️ Anh Hậu vui lòng tải đủ cả file 'Phụ lục 3' và 'Giáo án gốc' lên để AI đối chiếu nhé!")
